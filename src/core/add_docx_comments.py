import os
import json
import docx

def add_comment(paragraph, comment_text):
    """Insert a comment into the DOCX paragraph (simplified)."""
    run = paragraph.add_run(f" [COMMENT: {comment_text}]")
    run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)

def add_comments_to_docs(report_json_path, reviewed_folder):
    with open(report_json_path, "r", encoding="utf-8") as f:
        report = json.load(f)

    os.makedirs(reviewed_folder, exist_ok=True)

    for issue in report.get("issues_found", []):
        doc_name = issue.get("document")
        doc_type = issue.get("document_type", "Unknown")
        suggestion = issue.get("suggestion", "")

        raw_path = os.path.join("data", "raw", doc_name)
        if not os.path.exists(raw_path) or not doc_name.lower().endswith(".docx"):
            continue

        doc = docx.Document(raw_path)

        if issue.get("snippet"):
            snippet_lower = issue["snippet"].lower()
            for para in doc.paragraphs:
                if snippet_lower in para.text.lower():
                    add_comment(para, f"{doc_type}: {suggestion}")

        reviewed_path = os.path.join(reviewed_folder, doc_name)
        doc.save(reviewed_path)
        print(f" Reviewed file saved: {reviewed_path}")

    print(f" All documents processed and saved with comments in '{reviewed_folder}' folder.")

if __name__ == "__main__":
    REPORT_PATH = os.path.join("data", "reports", "report_issues.json")
    REVIEW_FOLDER = os.path.join("data", "reviewed")
    add_comments_to_docs(REPORT_PATH, REVIEW_FOLDER)
